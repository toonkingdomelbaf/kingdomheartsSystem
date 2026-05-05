import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import tkinter.font as tkFont
import unicodedata
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------- Helpers -----------------
def extract_plain_text(tk_text_widget):
    """
    Extract plain text from the Replace Text widget.
    """
    return tk_text_widget.get("1.0", "end-1c")


def normalize_text(s):
    """
    Normalize Word text to handle special characters like:
    - non-breaking hyphens
    - non-breaking spaces
    - Unicode normalization
    """
    s = s.replace("\u2011", "-").replace("\u00A0", " ")
    s = unicodedata.normalize("NFKC", s)
    return s


def get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf):
    """
    Return a list of files to process.
    """
    files = []
    if is_file:
        return [path]

    allowed_exts = []
    if txt: allowed_exts.append(".txt")
    if doc: allowed_exts.append(".docx")
    if pdf: allowed_exts.append(".pdf")

    for root, dirs, filenames in os.walk(path):
        for name in filenames:
            if not allowed_exts or any(name.lower().endswith(ext) for ext in allowed_exts):
                files.append(os.path.join(root, name))
        if not include_subfolders:
            break

    return files


# ----------------- TXT Replacement -----------------
def replace_in_file_txt(path, find_text, replace_text, case_sensitive, regex):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    flags = 0 if case_sensitive else re.IGNORECASE

    if regex:
        new_content = re.sub(find_text, replace_text, content, flags=flags)
    else:
        pattern = re.compile(re.escape(find_text), flags=flags)
        new_content = pattern.sub(replace_text, content)

    with open(path, "w", encoding="utf-8", errors="ignore") as f:
        f.write(new_content)



HIGHLIGHT_MAP = {
    0: "None", 1: "Black", 2: "Blue", 3: "Turquoise",
    4: "Bright Green", 5: "Pink", 6: "Red", 7: "Yellow", 8: "White",
    9: "Dark Blue", 10: "Teal", 11: "Green", 12: "Violet",
    13: "Dark Red", 14: "Dark Yellow", 15: "Gray50", 16: "Gray25"
}

WORD_TO_COLOR = {
    "Black": WD_COLOR_INDEX.BLACK, "Blue": WD_COLOR_INDEX.BLUE,
    "Turquoise": WD_COLOR_INDEX.TURQUOISE, "Bright Green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "Pink": WD_COLOR_INDEX.PINK, "Red": WD_COLOR_INDEX.RED, "Yellow": WD_COLOR_INDEX.YELLOW,
    "White": WD_COLOR_INDEX.WHITE, "Dark Blue": WD_COLOR_INDEX.DARK_BLUE, 
    "Teal": WD_COLOR_INDEX.TEAL, "Green": WD_COLOR_INDEX.GREEN, 
    "Violet": WD_COLOR_INDEX.VIOLET, "Dark Red": WD_COLOR_INDEX.DARK_RED,
    "Dark Yellow": WD_COLOR_INDEX.DARK_YELLOW, "Gray50": WD_COLOR_INDEX.GRAY_50, 
    "Gray25": WD_COLOR_INDEX.GRAY_25, "None": None
}

# Map Tkinter background colors to DOCX highlight colors if needed
BG_COLOR_TO_DOCX_HIGHLIGHT = {
    "#000000": WD_COLOR_INDEX.BLACK,
    "#0000FF": WD_COLOR_INDEX.BLUE,
    "#00FFFF": WD_COLOR_INDEX.TURQUOISE,
    "#00FF00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#FFC0CB": WD_COLOR_INDEX.PINK,
    "#FF0000": WD_COLOR_INDEX.RED,
    "#FFFF00": WD_COLOR_INDEX.YELLOW,
    "#FFFFFF": WD_COLOR_INDEX.WHITE,
    "#00008B": WD_COLOR_INDEX.DARK_BLUE,
    "#008080": WD_COLOR_INDEX.TEAL,
    "#008000": WD_COLOR_INDEX.GREEN,
    "#EE82EE": WD_COLOR_INDEX.VIOLET,
    "#8B0000": WD_COLOR_INDEX.DARK_RED,
    "#9B870C": WD_COLOR_INDEX.DARK_YELLOW,
    "#808080": WD_COLOR_INDEX.GRAY_50,
    "#C0C0C0": WD_COLOR_INDEX.GRAY_25,
}


def set_run_shading(run, hex_color):
    """
    Apply shading to a run using a hex color (e.g., #C0C0C0)
    """
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        rPr.append(shd)
    shd.set(qn('w:fill'), hex_color.replace('#',''))

def bg_to_docx_highlight(bg: str):
    """
    Convert Tkinter background hex color to DOCX highlight.
    Handles case-insensitive matching and fallback for common grey shades.
    """
    if not bg:
        return None
    bg = bg.upper()
    mapping = {
        "#000000": WD_COLOR_INDEX.BLACK,
        "#0000FF": WD_COLOR_INDEX.BLUE,
        "#00FFFF": WD_COLOR_INDEX.TURQUOISE,
        "#00FF00": WD_COLOR_INDEX.BRIGHT_GREEN,
        "#FFC0CB": WD_COLOR_INDEX.PINK,
        "#FF0000": WD_COLOR_INDEX.RED,
        "#FFFF00": WD_COLOR_INDEX.YELLOW,
        "#FFFFFF": WD_COLOR_INDEX.WHITE,
        "#00008B": WD_COLOR_INDEX.DARK_BLUE,
        "#008080": WD_COLOR_INDEX.TEAL,
        "#008000": WD_COLOR_INDEX.GREEN,
        "#EE82EE": WD_COLOR_INDEX.VIOLET,
        "#8B0000": WD_COLOR_INDEX.DARK_RED,
        "#9B870C": WD_COLOR_INDEX.DARK_YELLOW,
        "#808080": WD_COLOR_INDEX.GRAY_50,
        "#C0C0C0": WD_COLOR_INDEX.GRAY_25,
    }
    if bg in mapping:
        return mapping[bg]
    # fallback for other common grey shades
    if bg in ("#BEBEBE", "#D3D3D3"):
        return WD_COLOR_INDEX.GRAY_25
    return None


def get_text_widget_char_formats(text_widget):
    """
    Extract per-character text and formatting from a Tkinter Text widget.
    Returns a list of dictionaries for use in replace_paragraph_safe_inplace.
    """
    result = []
    text = text_widget.get("1.0", "end-1c")
    for i, ch in enumerate(text):
        idx = f"1.0 + {i} chars"
        # Default formatting
        font_family = "Arial"
        font_size = 12
        bold = False
        italic = False
        fg_color = "#000000"
        bg_color = None
        highlight_hex = None

        # Check all tags at this position
        tags = text_widget.tag_names(idx)
        for tag in tags:
            tag_config = text_widget.tag_cget(tag, "font")
            if tag_config:
                tk_font = tkFont.Font(font=text_widget.tag_cget(tag, "font"))
                font_family = tk_font.actual("family")
                font_size = tk_font.actual("size")
                bold = tk_font.actual("weight") == "bold"
                italic = tk_font.actual("slant") == "italic"

            fg = text_widget.tag_cget(tag, "foreground")
            if fg:
                fg_color = fg

            bg = text_widget.tag_cget(tag, "background")
            if bg:
                bg_color = bg_to_docx_highlight(bg)
                highlight_hex = bg  # store original hex for fallback

        # Convert fg_color #RRGGBB to RGBColor
        r = int(fg_color[1:3], 16)
        g = int(fg_color[3:5], 16)
        b = int(fg_color[5:7], 16)
        color_rgb = RGBColor(r, g, b)

        result.append({
            "text": ch,
            "font_family": font_family,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "color": color_rgb,
            "highlight": bg_color,
            "highlight_hex": highlight_hex
        })
    return result

# The function replace_paragraph_safe_inplace is responsible for replacing text within a Word paragraph (para) in-place 
# while preserving formatting and highlights.
def replace_paragraph_safe_inplace(para, pattern, char_formats):
    """
    Replaces only matched text in-place with per-character formatting.
    If char_formats is empty, removes the matched text.
    Unmatched text keeps all original formatting, including highlight and color.
    """
    if not para.text:
        return

    full_text = para.text
    matches = list(pattern.finditer(full_text))
    if not matches:
        return  # No matches, keep as is

    def force_run_font(run, font_name):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for k in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{k}"), font_name)
        run.font.name = font_name

    # Build mapping: run -> paragraph text positions
    runs = []
    pos = 0
    for run in para.runs:
        runs.append({
            "run": run,
            "start": pos,
            "end": pos + len(run.text),
            "font_name": run.font.name,
            "size": run.font.size,
            "bold": run.font.bold,
            "italic": run.font.italic,
            "color": run.font.color.rgb if run.font.color else None,
            "highlight": run.font.highlight_color
        })
        pos += len(run.text)

    new_runs = []
    cursor = 0

    for match in matches:
        start, end = match.span()

        # Text before match -> original formatting
        for r in runs:
            if r["end"] <= cursor or r["start"] >= start:
                continue
            s = max(cursor, r["start"])
            e = min(start, r["end"])
            if s < e:
                frag = r["run"].text[s - r["start"]: e - r["start"]]
                new_runs.append((frag, r, False))

        # Matched replacement
        if char_formats:
            for c in char_formats:
                new_runs.append((c["text"], c, True))

        cursor = end

    # Text after last match -> original formatting
    for r in runs:
        if r["end"] <= cursor:
            continue
        frag = r["run"].text[max(cursor - r["start"], 0):]
        if frag:
            new_runs.append((frag, r, False))

    # Clear paragraph and rebuild runs
    para.clear()
    for text, fmt, is_match in new_runs:
        r = para.add_run(text)
        if is_match:
            r.font.name = fmt["font_family"]
            r.font.size = Pt(fmt["font_size"])
            r.font.bold = fmt["bold"]
            r.font.italic = fmt["italic"]
            if fmt["color"]:
                r.font.color.rgb = fmt["color"]
            # Handle highlight
            if fmt.get("highlight") is not None:
                try:
                    r.font.highlight_color = fmt["highlight"]
                except Exception:
                    if fmt.get("highlight_hex"):
                        set_run_shading(r, fmt["highlight_hex"])
            elif fmt.get("highlight_hex"):
                # If highlight not in WD_COLOR_INDEX, still apply shading
                set_run_shading(r, fmt["highlight_hex"])
            force_run_font(r, fmt["font_family"])
        else:
            r.font.name = fmt["font_name"]
            r.font.size = fmt["size"]
            r.font.bold = fmt["bold"]
            r.font.italic = fmt["italic"]
            if fmt["color"]:
                r.font.color.rgb = fmt["color"]
            if fmt.get("highlight") is not None:
                try:
                    r.font.highlight_color = fmt["highlight"]
                except Exception:
                    if fmt.get("highlight_hex"):
                        set_run_shading(r, fmt["highlight_hex"])
            elif fmt.get("highlight_hex"):
                set_run_shading(r, fmt["highlight_hex"])

def process_tables(tables, handler):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    handler(para)
                if cell.tables:
                    process_tables(cell.tables, handler)

def replace_in_file_docx(path, tk_replace_widget, find_text,
                         case_sensitive=False, regex=False):
    """
    Replaces matched text in DOCX with Word selection formatting.
    If no text selected, removes the matched text.
    Only matched text is replaced; unmatched text keeps original formatting including highlight.
    """
    doc = Document(path)
    char_formats = get_text_widget_char_formats(tk_replace_widget)
    if char_formats is None:
        char_formats = []  # Treat as empty to remove matched text

    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(find_text, flags) if regex else re.compile(re.escape(find_text), flags)

    handler = lambda p: replace_paragraph_safe_inplace(p, pattern, char_formats)

    for para in doc.paragraphs:
        handler(para)
    process_tables(doc.tables, handler)

    doc.save(path)
    print(f"Replaced text in {path}.")
# end functions of replaced

# ----------------- Main Replace Process -----------------
def run_replace_process(gui, path, is_file, find_text, tk_replace_widget,
                        case_sensitive, regex, include_subfolders, txt, doc, pdf):
    
    """
    Main function to replace text in file/folder with formatting preserved.
    """
    gui.text_results.config(state="normal")
    files = get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf)
    gui.text_results.insert("end", f"Processing {len(files)} file(s)...\n")

    for file_path in files:
        try:
            if file_path.lower().endswith(".txt"):
                plain_text = extract_plain_text(tk_replace_widget)
                print("plain_text:" . plain_text + ":")
                replace_in_file_txt(file_path, find_text, plain_text, case_sensitive, regex)

            elif file_path.lower().endswith(".docx"):
                replace_in_file_docx(file_path, tk_replace_widget, find_text, case_sensitive, regex)
                # replace_in_file_docx(file_path, tk_replace_widget, find_text, case_sensitive, regex)

            elif file_path.lower().endswith(".pdf"):
                gui.text_results.insert("end", f"PDF replace not supported: {file_path}\n")
                continue

            gui.text_results.insert("end", f"Replaced in: {file_path}\n")

        except Exception as e:
            gui.text_results.insert("end", f"Error in {file_path}: {e}\n")

    gui.text_results.config(state="disabled")
